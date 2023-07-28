import re
import glob
from docx import Document

def split_and_save_text(file_path):
    # 读取.docx文件
    doc = Document(file_path)
    
    # 获取文档内容
    text = ''
    for para in doc.paragraphs:
        text += para.text
    
    # 使用正则表达式分割文本
    parts = re.split(r'(?<![一二三四五六七八九十])\s*[\r\n\s]*(?=[一二三四五六七八九十]、)', text)
    
    # 获取文件名（不含扩展名）
    file_name = file_path.split('\\')[-1].split('.')[0]
    print(file_name)
    
    # 遍历每个部分，并保存为.docx文件
    for i, part in enumerate(parts):
        # 创建一个新的Word文档
        new_doc = Document()
        
        # 添加部分内容到文档
        new_doc.add_paragraph(part)
        
        # 保存文档
        new_doc.save(f'./docx/output/{i}/{file_name}.docx')

# 获取docx文件夹下所有的.docx文件路径
file_paths = glob.glob('./docx/*.docx')

# 遍历每个文件执行分割和保存操作
# for file_path in file_paths:
#     split_and_save_text(file_path)

# Initialize variables
data = []
question = ""
answer = ""
question_started = False

file_path = 'C:\\Users\\20237\\Desktop\\2.docx'
doc = Document(file_path)

# Loop through each paragraph in the document
for para in doc.paragraphs:
    text = para.text.strip()

    # If the paragraph is a question
    if re.match(r'^\d+\.', text):
        question_started = True
        # If there is a previous question, add it to the data
        if question and answer:
            # find the position of the next question explanation in the answer
            next_question_index = answer.find("【{}题详解】".format(int(question.split('.')[0]) + 1))
            # if found, truncate the answer
            if next_question_index != -1:
                answer = answer[:next_question_index]
            data.append({"题目": question, "详解": answer.strip()})
        # Reset the question and answer
        question = text
        answer = ""
    # If the paragraph is an answer or a part of the question
    elif question_started:
        # If the next question starts
        if re.match(r'^\d+\.', text):
            question_started = False
            if question and answer:
                # find the position of the next question explanation in the answer
                next_question_index = answer.find("【{}题详解】".format(int(question.split('.')[0]) + 1))
                # if found, truncate the answer
                if next_question_index != -1:
                    answer = answer[:next_question_index]
                data.append({"题目": question, "详解": answer.strip()})
            # Start new question
            question = text
            answer = ""
        else:
            # If it's part of the question or answer
            if text.startswith("【{}题详解】".format(question.split('.')[0])) or not re.match(r'^\d+\.', text):
                answer += text + "\n"
            else:
                question += text + "\n"

# Add the last question to the data
if question and answer:
    # find the position of the next question explanation in the answer
    next_question_index = answer.find("【{}题详解】".format(int(question.split('.')[0]) + 1))
    # if found, truncate the answer
    if next_question_index != -1:
        answer = answer[:next_question_index]
    data.append({"题目": question, "详解": answer.strip()})

# Convert the data to JSON
json_data = json.dumps(data, ensure_ascii=False, indent=4)

json_data[:2000]  # Print first 2000 characters of the JSON data
with open("./questions_and_answers.json", "w", encoding="utf-8") as f:
    f.write(json_data)